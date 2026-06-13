"""
Regenera 'Architecture.docx' (español) — versión actualizada con la lógica nueva
de esta semana: borrado seguro "verificar y decidir" (VA02 + ME22N), idempotencia
a nivel de orden, deduplicación en Batch 8, matching resiliente de columnas,
observabilidad (logging a archivo), diagnóstico y manejo acotado de popups.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent / "Documentation" / "Architecture.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
GREY = RGBColor(0x59, 0x59, 0x59)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT
    return h


def para(text=""):
    return doc.add_paragraph(text)


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            r = p.add_run(it[0] + " — ")
            r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    return p


def table(headers, rows, style="Light Grid Accent 1"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = style
    for i, h in enumerate(headers):
        t.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].add_run(str(val))
    return t


# ===================== TÍTULO =====================
tp = doc.add_paragraph()
r = tp.add_run("Arquitectura — Proyecto DIL")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = ACCENT

sp = doc.add_paragraph()
r = sp.add_run("Pipeline autónomo de baja de tickets en SAP · Actualizado junio 2026")
r.italic = True
r.font.color.rgb = GREY

para()
p = doc.add_paragraph()
r = p.add_run("Este documento describe la arquitectura del sistema y la lógica nueva "
              "introducida esta semana (borrado seguro “verificar y decidir”, "
              "idempotencia a nivel de orden, deduplicación, matching resiliente de "
              "columnas, observabilidad con logging a archivo y herramienta de "
              "diagnóstico). Las secciones marcadas “NUEVO” reflejan los cambios de "
              "esta semana.")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = GREY

# ===================== OVERVIEW =====================
heading("Visión general", 1)
para(
    "El pipeline automatiza un proceso que antes era manual en SAP ECC (producción): "
    "tomar una lista de tickets de transporte y eliminar/revertir todos sus documentos "
    "asociados en el orden correcto. Piénsalo como una “cadena de desmontaje” — cada "
    "pieza debe quitarse antes de poder quitar la siguiente. El bot maneja el SAP GUI "
    "real vía SAP GUI Scripting (COM), corre tres sesiones SAP en paralelo, procesa "
    "los tickets en chunks resilientes, verifica cada paso contra el propio SAP, se "
    "recupera de interrupciones y emite un reporte auditable."
)

# ===================== 1. SESIONES =====================
heading("1. Sesiones SAP", 1)
para("El pipeline abre tres sesiones SAP simultáneas desde la misma conexión, cada una "
     "con una responsabilidad exclusiva. Esto evita que las eliminaciones interrumpan "
     "las lecturas de VL06F.")
table(
    ["Sesión", "Responsabilidad", "Transacciones", "Cuándo actúa"],
    [
        ["session1 — Lectura/Verificación", "Leer datos y verificar; también ejecuta BOL",
         "VL06F", "Al inicio (lectura) y después de cada batch (verificación)"],
        ["session2 — Operaciones", "Eliminar documentos en el orden correcto",
         "VF11 · VI05 · VT02N · VL09", "En cada batch; nunca navega a VL06F"],
        ["session3 — Órdenes", "Leer ZCMR y eliminar/cancelar órdenes",
         "ZCMR · VA02 · ME22N", "Batch 6 (inicial) y Batch 8 (retry)"],
    ],
)
para()
p = doc.add_paragraph()
r = p.add_run("Fallback: ")
r.bold = True
p.add_run("si SAP no permite sesiones adicionales, sap_login.py asigna "
          "session2 = session3 = session1 y el pipeline continúa con una sola sesión.")

# ===================== 2. FLUJO =====================
heading("2. Flujo completo (chunked)", 1)
para("Los tickets se parten en chunks de 100 (configurable). Cada chunk se procesa "
     "completo por todos los batches antes de pasar al siguiente. Las sesiones SAP se "
     "reúsan entre chunks (no se vuelve a logear).")
code(
    "POR cada chunk de ~100 tickets:\n"
    "  4.1  LEER VL06F           [session1]  una consulta -> todos los docs en memoria\n"
    "  4.2  BATCH 1  VF11        [session2]  reversar billing docs (bulk)   + verificar\n"
    "  4.3  BATCH 2  VI05        [session2]  eliminar shipment costs (estado C) + verificar\n"
    "  4.4  BATCH 3  VT02N       [session2]  desasignar + eliminar shipment   + verificar\n"
    "  4.5  BATCH 4  VL09        [session2]  reversar PGI (bulk, WBSTK != 'A') + verificar\n"
    "  4.6  BATCH 5  VL06F       [session1]  eliminar BOL / delivery (por ticket)\n"
    "  4.7  BATCH 6  ZCMR->VA02/ME22N [s3]   eliminar órdenes (línea por línea)\n"
    "  4.8  BATCH 7  ZSD_DEL_TICKETS  [s1]   verificación contra ground truth\n"
    "  4.9  BATCH 8  VA02/ME22N   [s3]       retry pendientes (dedup + verificar-y-decidir)\n"
    "  4.10 ACUMULAR resultados del chunk al diccionario consolidado\n"
    "AL TERMINAR todos los chunks:  reporte consolidado (consola + xlsx)"
)
para("Una sola consulta a VL06F carga en memoria todos los números de documento del "
     "chunk; todos los batches leen de ese snapshot en lugar de re-consultar SAP — una "
     "optimización clave.")

# ===================== 3. CHUNKING =====================
heading("3. Por qué chunking", 1)
para("SAP VL09 (reversar PGI) tiene un límite práctico de ~60–100 tickets por bulk:")
bullets([
    "Timeout del work process de diálogo (rdisp/max_wprun_time = 600 s por defecto)",
    "Desbordamiento de la cola de update task",
    "Conflictos de lock entre múltiples deliveries",
])
para("Chunks de 100 mantienen cada operación en el rango seguro y reducen el blast "
     "radius de un fallo: si un chunk falla por timeout, los demás siguen.")

# ===================== 4. CHECKPOINTING =====================
heading("4. Checkpointing y resume", 1)
para("El progreso se persiste por chunk a disco para sobrevivir interrupciones "
     "(timeout, kill, crash, reinicio). Al re-correr salta los chunks completados.")
code(
    "Data-bases/Estado/{nombre_excel}_state.json\n"
    "  version, run_id, started_at\n"
    "  input_file, input_file_hash (SHA256)\n"
    "  chunk_size, total_tickets, total_chunks, batches_to_run\n"
    "  chunks: { \"0\": {tickets, resultados, vl06f, duración, ...}, \"1\": {...} }"
)
table(
    ["Evento", "Comportamiento"],
    [
        ["Primera corrida", "Crea el state con el hash del Excel; agrega chunks al completar"],
        ["Re-corrida, Excel sin cambios", "Carga state, salta chunks completados, continúa"],
        ["Re-corrida, Excel modificado", "Hash mismatch -> descarta state, arranca fresh"],
        ["Flag --fresh", "Elimina el state file antes de arrancar"],
        ["Flag --report-only", "No toca SAP; consolida el state y emite xlsx"],
        ["Crash mid-chunk", "El chunk no se guarda -> se reintenta completo al re-correr"],
        ["Escritura del state", "Atómica (.json.tmp + rename) -> nunca queda corrupto"],
    ],
)

# ===================== 5. BORRADO SEGURO (NUEVO) =====================
heading("5. Borrado seguro de órdenes: “verificar y decidir”  (NUEVO)", 1)
para(
    "El cambio central de esta semana. En este SAP una misma orden mezcla con "
    "frecuencia líneas que se deben borrar con líneas que se deben conservar. El "
    "Batch 8 anterior borraba la orden COMPLETA vía menú, lo que destruiría "
    "silenciosamente datos que deben permanecer. El motor nuevo inspecciona el estado "
    "real de cada orden y decide, línea por línea, qué es seguro borrar."
)
p = doc.add_paragraph()
p.add_run("Punto de entrada unificado:").bold = True
code("_delete_order_lines_for_tickets(session, order, tickets)   # tickets = set en alcance")
para("Despacha según el tipo de orden y aplica la misma lógica a ambas:")
table(
    ["Tipo de orden", "Detección", "Transacción", "Función"],
    [
        ["Sales order intercompany", "default", "VA02", "_delete_intercompany_lines_va02"],
        ["Purchase order intracompany", "order[:2] == '47'", "ME22N", "_delete_intracompany_lines_me22n"],
    ],
)
para()
p = doc.add_paragraph()
p.add_run("Lógica de decisión (por orden):").bold = True
bullets([
    ("Leer", "abrir la orden una vez y leer sus líneas (PO Number BSTKD_E = ticket, con "
     "ceros a la izquierda). Lectura y selección en una sola pasada, porque una fila "
     "solo se puede seleccionar mientras está visible en pantalla."),
    ("Clasificar", "comparar cada línea contra el set de tickets en alcance (padded): en "
     "alcance vs. a conservar."),
    ("Todas en alcance", "borrar todas las líneas seleccionadas vacía la orden y SAP la "
     "elimina sola."),
    ("Orden mezclada", "seleccionar y borrar SOLO las líneas en alcance; el resto persiste."),
    ("Ninguna en alcance", "las líneas ya no están -> se cuenta éxito (idempotente)."),
    ("Fail-safe", "si una línea no se puede clasificar con certeza, el bot se detiene en "
     "vez de adivinar — nunca borra datos de los que no está seguro. ME22N además vuelca "
     "sus nombres de columna reales al log la primera vez, para confirmar el campo del ticket."),
])
para("Esto garantiza que las líneas que deben conservarse nunca se borran — una "
     "propiedad que el proceso manual no podía prometer con fiabilidad.")

# ===================== 6. IDEMPOTENCIA (NUEVO) =====================
heading("6. Idempotencia a nivel de orden  (NUEVO)", 1)
para("Re-correr el pipeline nunca debe fallar en órdenes que ya fueron borradas en una "
     "corrida previa. Cuando VA02/ME22N reporta que la orden ya no existe, se cuenta "
     "como éxito, no como fallo.")
code(
    "_order_already_gone(session, order, popup_msg)  ->  True si SAP dice que ya no está\n"
    "_ORDER_GONE_KEYWORDS: 'is not in the database', 'has been archived',\n"
    "                     'does not exist', 'no existe', ...\n"
    "Mensaje real de SAP confirmado en producción:\n"
    "  'SD document 1150780411 is not in the database or has been archived'"
)
para("Si la orden no abre en modo edición por otra razón, el bot lanza un error claro "
     "con el texto real del status bar de SAP — en lugar del críptico “control could "
     "not be found by id”.")

# ===================== 7. DEDUP (NUEVO) =====================
heading("7. Deduplicación de órdenes en Batch 8  (NUEVO)", 1)
para("Decenas de tickets suelen apuntar a la misma orden. Antes, Batch 8 intentaba "
     "borrar la misma orden una vez por ticket (ej. una orden 50 veces). "
     "_cancel_orders_deduped ahora agrupa orden -> [tickets], procesa cada orden "
     "distinta UNA vez y propaga el resultado a todos sus tickets.")
code(
    "ticket_to_order (150 tickets)  ->  agrupar por orden  ->  ~4 órdenes distintas\n"
    "por cada orden distinta:\n"
    "    _delete_order_lines_for_tickets(session, order, set(tickets_de_esa_orden))\n"
    "    el éxito/fallo se aplica a todos los tickets de la orden"
)
para("Nota: el flujo manual línea-por-línea (cancel_order_by_ticket) NO se deduplica a "
     "propósito — ahí cada ticket es una línea distinta.")

# ===================== 8. COLUMNAS RESILIENTES (NUEVO) =====================
heading("8. Matching resiliente de columnas de grid  (NUEVO)", 1)
para("SAP expone nombres técnicos de columna distintos según el layout activo, lo que "
     "hacía que las lecturas de grid devolvieran nada en silencio. El pipeline ahora "
     "resuelve columnas contra una lista de candidatos y vuelca al log los nombres "
     "reales la primera vez que abre un grid.")
bullets([
    ("_resolve_col / _dump_subgrid_columns", "sub-grid de ZCMR (columnas de orden/"
     "delivery/ticket). Nombres reales confirmados: TICKET_CODE, SD_ORDER, DELIVERY."),
    ("verify_zsd_del_tickets", "grid de ground truth ZSD: la columna del ticket es "
     "TICKET_CODE (no TICKET) — corregido."),
])

# ===================== 9. OBSERVABILIDAD (NUEVO) =====================
heading("9. Observabilidad — logging a archivo  (NUEVO)", 1)
para("Cada corrida duplica toda la salida de consola a un log con timestamp en UTF-8, "
     "dando una traza completa y auditable de cada decisión y mensaje de SAP. "
     "Implementado de forma central en log_util.py vía un Tee que escribe a consola y "
     "archivo a la vez, robusto al encoding de la consola de Windows (cp1252) para que "
     "ninguna línea se pierda en silencio.")
code("Data-bases/Logs/dil_run_YYYYMMDD_HHMMSS.log     (UTF-8, traza completa)")

# ===================== 10. DIAGNÓSTICO (NUEVO) =====================
heading("10. Herramienta de diagnóstico  (NUEVO)", 1)
para("diagnose_grids.py se conecta a una sesión SAP viva y vuelca los nombres técnicos "
     "reales de columna y los layouts disponibles de ZCMR, ZSD y VL06F, sin modificar "
     "nada (solo lectura). Este enfoque de “instrumentar primero, arreglar después” es "
     "cómo se diagnosticaron problemas sutiles y de alto riesgo desde una sola corrida.")
code("python diagnose_grids.py zcmr <ticket>     # también: zsd | vl06f | all | layouts")

# ===================== 11. ROBUSTEZ POPUPS (NUEVO) =====================
heading("11. Robustez — manejo acotado de popups  (NUEVO)", 1)
para("VA02 muestra una cascada de popups después de guardar el borrado de una línea. "
     "El handler tiene ahora tres cotas de seguridad para que un popup que no se cierra "
     "nunca cuelgue el pipeline: tope de popups, deadline absoluto (que no se resetea) y "
     "detección del mismo popup reapareciendo. Si se dispara cualquiera, la orden se "
     "marca como fallida (revisable) en vez de quedar en bucle infinito.")

# ===================== 12. EJECUCIÓN PARCIAL =====================
heading("12. Ejecución parcial del pipeline", 1)
table(
    ["Flag", "Default", "Descripción"],
    [
        ["--batches", "all", "Subset de batches: '1,2,3' o '6-8' o '1,4,6-8'"],
        ["--chunk-size", "100", "Tickets por chunk (límite práctico VL09: 60–100)"],
        ["--fresh", "false", "Elimina el state previo antes de arrancar"],
        ["--report-only", "false", "No login SAP; consolida state y emite xlsx"],
        ["--manual-only", "false", "Recovery: procesa pares (ticket, order) de la hoja "
         "'Manual Orders' línea-por-línea (solo Batches 6–8)"],
        ["--retry-failed", "false", "Reintenta solo los tickets que quedaron fallidos en el state"],
    ],
)
para("Si no se selecciona ningún batch 1–5, el pipeline omite la lectura de VL06F y pasa "
     "todos los tickets directo a Batch 6 — útil para re-procesar órdenes pendientes "
     "cuando los tickets ya se eliminaron de VL06F.")

# ===================== 13. REGLAS DEL FLUJO =====================
heading("13. Reglas del flujo", 1)
bullets([
    "Los tickets son independientes — el fallo de uno nunca detiene a los demás.",
    "El orden de los batches es obligatorio — SAP bloquea pasos posteriores si quedan docs previos.",
    "Cada batch opera solo sobre los tickets que pasaron el batch anterior.",
    "Si una operación bulk falla, la verificación VL06F igual clasifica correctamente "
    "los tickets afectados como fallidos.",
    "Si la verificación falla, los tickets quedan como fallidos por seguridad — nunca falso positivo.",
])

# ===================== 14. ESTRUCTURA DE ARCHIVOS =====================
heading("14. Estructura de archivos", 1)
code(
    "Proyecto DIL/\n"
    "  code/\n"
    "    main.py            # orquestador: chunks + checkpoint + dedup Batch 8 + reporte\n"
    "    sap_login.py       # login + abre las 3 sesiones SAP\n"
    "    sap_utils.py       # helpers compartidos (_navigate_to, _wait_ready, multi-valor)\n"
    "    sap_vl06f.py       # session1: read_vl06f_data, delete_bol\n"
    "    sap_batches.py     # session2: VF11, VI05, VT02N, VL09\n"
    "    sap_orders.py      # session3: ZCMR, VA02, ME22N + motor verificar-y-decidir\n"
    "    verifications.py   # session1: verificación bulk post-batch (+ ground truth ZSD)\n"
    "    checkpoint.py      # state JSON por chunk (resume)\n"
    "    excel_reader.py    # lee tickets (y la hoja 'Manual Orders')\n"
    "    report_writer.py   # reporte xlsx final\n"
    "    path.py            # resuelve rutas Entradas/Salidas/Estado\n"
    "    log_util.py        # NUEVO: logging consola+archivo (Tee, seguro a encoding)\n"
    "    diagnose_grids.py  # NUEVO: diagnóstico read-only de grids/layouts de SAP\n"
    "  Data-bases/\n"
    "    Entradas/<mes>/    # Excel de entrada\n"
    "    Salidas/<mes>/     # reporte xlsx generado\n"
    "    Estado/            # state files JSON (resume)\n"
    "    Logs/              # NUEVO: logs por corrida (UTF-8)\n"
    "  Documentation/"
)

# ===================== 15. FUNCIONES CLAVE =====================
heading("15. Funciones clave — sap_orders.py (session3)", 1)
table(
    ["Función", "Rol"],
    [
        ["_delete_order_lines_for_tickets", "Entry point seguro unificado; despacha VA02 vs ME22N"],
        ["_delete_intercompany_lines_va02", "VA02 verificar-y-decidir (leer+seleccionar en una pasada)"],
        ["_delete_intracompany_lines_me22n", "ME22N verificar-y-decidir (instrumenta primero)"],
        ["_order_already_gone / _read_sbar", "Idempotencia: orden ya borrada = éxito"],
        ["_resolve_col / _dump_subgrid_columns", "Matching resiliente del sub-grid de ZCMR"],
        ["_handle_va02_post_save_popups", "Manejo acotado de popups (anti-bucle)"],
    ],
)
para()
p = doc.add_paragraph()
p.add_run("Helper de deduplicación en main.py: ").bold = True
p.add_run("_cancel_orders_deduped(session, ticket_to_order, label, cancel_failures) "
          "agrupa órdenes y llama al entry point seguro una vez por orden distinta.")

# ===================== 16. REPORTE FINAL =====================
heading("16. Reporte final", 1)
para("Se genera una sola vez, al terminar todos los chunks, desde el state consolidado. "
     "Lista, por batch, los tickets exitosos y fallidos, y un resumen final de tickets "
     "completados al 100% vs. los que requieren revisión manual. Batch 8 solo procesa "
     "los tickets que Batch 7 mostró pendientes en ZSD_DEL_TICKETS.")
code(
    "===================================================\n"
    "  REPORTE CONSOLIDADO - TODOS LOS CHUNKS\n"
    "===================================================\n"
    "  BATCH 6 - ZCMR Orders     Exitosos: ...   Fallidos: ...\n"
    "  BATCH 7 - Final ZSD       Exitosos: ...   Fallidos: ...\n"
    "  BATCH 8 - Order Cancel    Exitosos: ...   Fallidos: ...\n"
    "  ---------------------------------------------------\n"
    "  Completados 100%:  N/total      Revisar manualmente:  ...\n"
    "==================================================="
)

doc.add_paragraph().add_run("─" * 60).font.color.rgb = ACCENT
foot = doc.add_paragraph()
r = foot.add_run("Validado en producción (junio 2026): un lote de 200 tickets se "
                 "decomisionó de punta a punta con éxito, sin intervención manual.")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = GREEN

doc.save(str(OUT))
print(f"OK -> {OUT}")
