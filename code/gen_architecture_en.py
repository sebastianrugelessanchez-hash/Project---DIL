"""
Genera 'Architecture (EN).docx' — versión en inglés del documento de arquitectura
del Proyecto DIL, actualizada con la lógica nueva de esta semana:
  - Borrado seguro "verify-then-decide" (VA02 + ME22N)
  - Idempotencia a nivel de orden
  - Deduplicación de órdenes en Batch 8
  - Matching resiliente de columnas de grid
  - Observabilidad (logging a archivo)
  - Herramienta de diagnóstico (diagnose_grids.py)
  - Manejo acotado de popups (anti-bucle)
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parent.parent / "Documentation" / "Architecture (EN).docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
GREY = RGBColor(0x59, 0x59, 0x59)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT
    return h


def para(text=""):
    return doc.add_paragraph(text)


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            r = p.add_run(it[0] + " — ")
            r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    return p


def newflag():
    """Inline 'NEW (this week)' marker run inside a heading paragraph."""
    pass


def table(headers, rows, style="Light Grid Accent 1"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = style
    for i, h in enumerate(headers):
        t.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].add_run(str(val))
    return t


# ======================= TITLE =======================
tp = doc.add_paragraph()
r = tp.add_run("Architecture — Project DIL")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = ACCENT

sp = doc.add_paragraph()
r = sp.add_run("Autonomous SAP ticket-decommissioning pipeline · Updated June 2026")
r.italic = True
r.font.color.rgb = GREY

para()
p = doc.add_paragraph()
r = p.add_run("This document describes the system architecture and the new logic "
              "introduced this week (safe “verify-then-decide” order deletion, "
              "order-level idempotency, de-duplication, resilient grid-column "
              "matching, file-based observability, and diagnostics tooling). "
              "Sections marked “NEW” reflect changes made this week.")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = GREY

# ======================= OVERVIEW =======================
heading("Overview", 1)
para(
    "The pipeline automates a process that used to be performed manually in SAP ECC "
    "(production): take a list of transport tickets and delete / reverse all of their "
    "associated documents in the correct order. Think of it as a controlled "
    "“teardown chain” — each piece must be removed before the next one can be removed. "
    "The bot drives the real SAP GUI through SAP GUI Scripting (COM), runs three "
    "parallel SAP sessions, processes tickets in resilient chunks, verifies every "
    "step against SAP itself, recovers from interruptions, and emits an auditable "
    "report."
)

# ======================= 1. SESSIONS =======================
heading("1. SAP sessions", 1)
para("The pipeline opens three simultaneous SAP sessions from the same connection, "
     "each with an exclusive responsibility. This keeps deletion operations from "
     "interrupting the VL06F reads.")
table(
    ["Session", "Responsibility", "Transactions", "When it acts"],
    [
        ["session1 — Read/Verify", "Read data and verify results; also runs BOL deletion",
         "VL06F", "At start (read) and after every batch (verify)"],
        ["session2 — Operations", "Delete documents in the correct order",
         "VF11 · VI05 · VT02N · VL09", "In every batch; never navigates to VL06F"],
        ["session3 — Orders", "Read ZCMR and delete/cancel orders",
         "ZCMR · VA02 · ME22N", "Batch 6 (initial) and Batch 8 (retry)"],
    ],
)
para()
p = doc.add_paragraph()
r = p.add_run("Fallback: ")
r.bold = True
p.add_run("if SAP does not allow additional sessions, sap_login.py assigns "
          "session2 = session3 = session1 and the pipeline continues with a single "
          "session.")

# ======================= 2. FLOW =======================
heading("2. End-to-end flow (chunked)", 1)
para("Tickets are split into chunks of 100 (configurable). Each chunk is processed "
     "completely by all batches before moving on. SAP sessions are reused across "
     "chunks (no re-login).")
code(
    "FOR each chunk of ~100 tickets:\n"
    "  4.1  READ VL06F           [session1]  one query -> all document numbers in memory\n"
    "  4.2  BATCH 1  VF11        [session2]  bulk-reverse billing documents   + verify\n"
    "  4.3  BATCH 2  VI05        [session2]  delete shipment costs (unlock C)  + verify\n"
    "  4.4  BATCH 3  VT02N       [session2]  un-assign + delete shipment       + verify\n"
    "  4.5  BATCH 4  VL09        [session2]  bulk-reverse PGI (WBSTK != 'A')   + verify\n"
    "  4.6  BATCH 5  VL06F       [session1]  delete BOL / delivery (per ticket)\n"
    "  4.7  BATCH 6  ZCMR->VA02/ME22N [s3]   delete orders (line-level, per ticket)\n"
    "  4.8  BATCH 7  ZSD_DEL_TICKETS  [s1]   ground-truth verification\n"
    "  4.9  BATCH 8  VA02/ME22N   [s3]       retry pending orders (de-dup, verify-then-decide)\n"
    "  4.10 ACCUMULATE chunk results into the consolidated dictionary\n"
    "AFTER all chunks:  consolidated report (console + xlsx)"
)
para("A single VL06F query loads every document number for the chunk into memory; all "
     "batches read from that snapshot instead of re-querying SAP — a key optimization.")

# ======================= 3. WHY CHUNKING =======================
heading("3. Why chunking", 1)
para("SAP VL09 (reverse PGI) has a practical limit of ~60–100 tickets per bulk run:")
bullets([
    "Dialog work-process timeout (rdisp/max_wprun_time = 600 s default)",
    "Update task queue overflow",
    "Lock conflicts between multiple deliveries",
])
para("Chunks of 100 keep each operation within the safe range and reduce the blast "
     "radius of a failure: if one chunk fails on a timeout, the others continue.")

# ======================= 4. CHECKPOINTING =======================
heading("4. Checkpointing & resume", 1)
para("Progress is persisted per chunk to disk so the pipeline survives interruptions "
     "(timeout, kill, crash, reboot). On re-run it skips completed chunks.")
code(
    "Data-bases/Estado/{excel_name}_state.json\n"
    "  version, run_id, started_at\n"
    "  input_file, input_file_hash (SHA256)\n"
    "  chunk_size, total_tickets, total_chunks, batches_to_run\n"
    "  chunks: { \"0\": {tickets, resultados, vl06f, duration, ...}, \"1\": {...} }"
)
table(
    ["Event", "Behaviour"],
    [
        ["First run", "Creates state with the Excel hash; appends chunks as they complete"],
        ["Re-run, Excel unchanged", "Loads state, skips completed chunks, continues"],
        ["Re-run, Excel modified", "Hash mismatch -> discards state, starts fresh"],
        ["--fresh flag", "Deletes the state file before starting"],
        ["--report-only flag", "Does not touch SAP; consolidates state and emits xlsx"],
        ["Crash mid-chunk", "Chunk is not saved -> retried in full on re-run"],
        ["State write", "Atomic (.json.tmp + rename) -> never left corrupt"],
    ],
)

# ======================= 5. SAFE ORDER DELETION (NEW) =======================
h = heading("5. Safe order deletion: “verify-then-decide”  (NEW)", 1)
para(
    "This week's central change. In this SAP system a single order frequently mixes "
    "lines that must be deleted with lines that must be kept. The previous Batch 8 "
    "deleted the WHOLE order via the menu, which would silently destroy data that has "
    "to be retained. The new engine inspects each order's real state and decides, "
    "per line, what is safe to delete."
)
para("Unified entry point: ").runs[0].bold = True
code("_delete_order_lines_for_tickets(session, order, tickets)   # tickets = in-scope set")
para("It dispatches by order type and applies the same decision logic to both:")
table(
    ["Order type", "Detection", "Transaction", "Function"],
    [
        ["Intercompany sales order", "default", "VA02", "_delete_intercompany_lines_va02"],
        ["Intracompany purchase order", "order[:2] == '47'", "ME22N", "_delete_intracompany_lines_me22n"],
    ],
)
para()
p = doc.add_paragraph()
r = p.add_run("Decision logic (per order): ")
r.bold = True
bullets([
    ("Read", "open the order once and read its line items (PO Number BSTKD_E = ticket, "
     "zero-padded). Read and select happen in a single pass, because a row can only be "
     "selected while it is visible on screen."),
    ("Classify", "match every line against the padded in-scope ticket set: in-scope vs. "
     "to-keep."),
    ("All lines in scope", "deleting all selected lines empties the order and SAP "
     "removes it."),
    ("Mixed order", "select and delete ONLY the in-scope lines; the rest persist."),
    ("No in-scope lines", "the lines are already gone -> treated as success (idempotent)."),
    ("Fail safe", "if a line cannot be positively classified, the bot stops rather than "
     "guessing — it never deletes data it is unsure about. ME22N additionally dumps its "
     "real column names to the log on first use so the ticket field can be confirmed."),
])
para("This guarantees that lines which must be kept are never deleted — a property the "
     "manual process could not reliably promise.")

# ======================= 6. ORDER-LEVEL IDEMPOTENCY (NEW) =======================
heading("6. Order-level idempotency  (NEW)", 1)
para("Re-running the pipeline must never error on orders that were already deleted in a "
     "previous run. When VA02/ME22N reports that an order is gone, it is counted as "
     "success, not failure.")
code(
    "_order_already_gone(session, order, popup_msg)  ->  True if SAP says it's gone\n"
    "_ORDER_GONE_KEYWORDS: 'is not in the database', 'has been archived',\n"
    "                     'does not exist', 'no existe', ...\n"
    "Real SAP message confirmed in production:\n"
    "  'SD document 1150780411 is not in the database or has been archived'"
)
para("If the order does not open in edit mode for any other reason, the bot raises a "
     "clear error containing SAP's real status-bar text — instead of a cryptic "
     "“control could not be found by id”.")

# ======================= 7. DE-DUPLICATION (NEW) =======================
heading("7. Order de-duplication in Batch 8  (NEW)", 1)
para("Dozens of tickets often map to the same order. Previously Batch 8 attempted to "
     "delete the same order once per ticket (e.g. one order tried 50 times). "
     "_cancel_orders_deduped now groups orders -> [tickets], processes each distinct "
     "order once, and propagates the result to all of its tickets.")
code(
    "ticket_to_order (150 tickets)  ->  grouped by order  ->  ~4 distinct orders\n"
    "for each distinct order:\n"
    "    _delete_order_lines_for_tickets(session, order, set(tickets_of_that_order))\n"
    "    success/failure is applied to every ticket sharing the order"
)
para("Note: the line-level manual flow (cancel_order_by_ticket) is intentionally NOT "
     "de-duplicated — there each ticket is a different line.")

# ======================= 8. RESILIENT COLUMN MATCHING (NEW) =======================
heading("8. Resilient grid-column matching  (NEW)", 1)
para("SAP exposes different technical column names depending on the active screen "
     "layout, which caused grid reads to silently return nothing. The pipeline now "
     "resolves columns against a list of candidate names and dumps the real column "
     "names to the log the first time a grid is opened.")
bullets([
    ("_resolve_col / _dump_subgrid_columns", "ZCMR sub-grid (order/delivery/ticket "
     "columns). Confirmed real names: TICKET_CODE, SD_ORDER, DELIVERY."),
    ("verify_zsd_del_tickets", "ZSD ground-truth grid: the ticket column is TICKET_CODE "
     "(not TICKET) — fixed."),
])

# ======================= 9. OBSERVABILITY (NEW) =======================
heading("9. Observability — file logging  (NEW)", 1)
para("Every run mirrors all console output to a timestamped UTF-8 log, giving a full, "
     "auditable trace of every decision and SAP message. Implemented centrally in "
     "log_util.py via a Tee that writes to console and file at once and is robust to "
     "Windows console encoding (cp1252) so no line is ever silently dropped.")
code("Data-bases/Logs/dil_run_YYYYMMDD_HHMMSS.log     (UTF-8, full trace)")

# ======================= 10. DIAGNOSTICS (NEW) =======================
heading("10. Diagnostics tooling  (NEW)", 1)
para("diagnose_grids.py connects to a live SAP session and dumps the real technical "
     "column names and available layouts for ZCMR, ZSD and VL06F, without modifying "
     "anything (read-only). This “instrument first, then fix” approach is how subtle, "
     "high-stakes issues were diagnosed from a single production run.")
code("python diagnose_grids.py zcmr <ticket>     # also: zsd | vl06f | all | layouts")

# ======================= 11. ROBUSTNESS: POPUPS (NEW) =======================
heading("11. Robustness — bounded popup handling  (NEW)", 1)
para("VA02 shows a cascade of popups after saving a line deletion. The handler now has "
     "three safety bounds so a popup that fails to close can never hang the pipeline: a "
     "maximum popup count, an absolute (non-resetting) deadline, and detection of the "
     "same popup reappearing. If any bound trips, the order is failed (and reviewable) "
     "instead of looping forever.")

# ======================= 12. PARTIAL EXECUTION =======================
heading("12. Partial pipeline execution", 1)
table(
    ["Flag", "Default", "Description"],
    [
        ["--batches", "all", "Subset of batches: '1,2,3' or '6-8' or '1,4,6-8'"],
        ["--chunk-size", "100", "Tickets per chunk (VL09 practical limit 60–100)"],
        ["--fresh", "false", "Delete previous state before starting"],
        ["--report-only", "false", "No SAP login; consolidate state and emit xlsx"],
        ["--manual-only", "false", "Recovery: process (ticket, order) pairs from the "
         "'Manual Orders' sheet, line-by-line (Batches 6–8 only)"],
        ["--retry-failed", "false", "Retry only the tickets left as failed in the state"],
    ],
)
para("If no batch 1–5 is selected, the pipeline skips the VL06F read and passes all "
     "tickets straight to Batch 6 — useful to re-process pending orders after the "
     "tickets were already removed from VL06F.")

# ======================= 13. FLOW RULES =======================
heading("13. Flow rules", 1)
bullets([
    "Tickets are independent — one ticket's failure never stops the others.",
    "Batch order is mandatory — SAP blocks later steps if earlier docs remain.",
    "Each batch operates only on tickets that passed the previous batch.",
    "If a bulk operation fails, the VL06F verification still classifies affected "
    "tickets correctly as failed.",
    "If verification fails, tickets are marked failed for safety — never a false positive.",
])

# ======================= 14. FILE STRUCTURE =======================
heading("14. File structure", 1)
code(
    "Proyecto DIL/\n"
    "  code/\n"
    "    main.py            # orchestrator: chunks + checkpoint + dedup Batch 8 + report\n"
    "    sap_login.py       # login + opens the 3 SAP sessions\n"
    "    sap_utils.py       # shared helpers (_navigate_to, _wait_ready, multi-value)\n"
    "    sap_vl06f.py       # session1: read_vl06f_data, delete_bol\n"
    "    sap_batches.py     # session2: VF11, VI05, VT02N, VL09\n"
    "    sap_orders.py      # session3: ZCMR, VA02, ME22N + verify-then-decide engine\n"
    "    verifications.py   # session1: bulk post-batch verification (+ ZSD ground truth)\n"
    "    checkpoint.py      # per-chunk JSON state (resume)\n"
    "    excel_reader.py    # reads tickets (and the 'Manual Orders' sheet)\n"
    "    report_writer.py   # final xlsx report\n"
    "    path.py            # resolves Entradas/Salidas/Estado paths\n"
    "    log_util.py        # NEW: console+file logging (Tee, encoding-safe)\n"
    "    diagnose_grids.py  # NEW: read-only SAP grid/layout diagnostics\n"
    "  Data-bases/\n"
    "    Entradas/<month>/  # input Excel\n"
    "    Salidas/<month>/   # generated xlsx report\n"
    "    Estado/            # JSON state files (resume)\n"
    "    Logs/              # NEW: per-run UTF-8 logs\n"
    "  Documentation/"
)

# ======================= 15. KEY MODULES (orders) =======================
heading("15. Key functions — sap_orders.py (session3)", 1)
table(
    ["Function", "Role"],
    [
        ["_delete_order_lines_for_tickets", "Unified safe entry point; dispatches VA02 vs ME22N"],
        ["_delete_intercompany_lines_va02", "VA02 verify-then-decide (read+select in one pass)"],
        ["_delete_intracompany_lines_me22n", "ME22N verify-then-decide (instrument-first)"],
        ["_order_already_gone / _read_sbar", "Idempotency: already-deleted order = success"],
        ["_resolve_col / _dump_subgrid_columns", "Resilient ZCMR sub-grid column matching"],
        ["_handle_va02_post_save_popups", "Bounded popup handling (anti-infinite-loop)"],
    ],
)
para()
p = doc.add_paragraph()
r = p.add_run("main.py de-duplication helper: ")
r.bold = True
p.add_run("_cancel_orders_deduped(session, ticket_to_order, label, cancel_failures) "
          "groups orders and calls the safe entry point once per distinct order.")

# ======================= 16. FINAL REPORT =======================
heading("16. Final report", 1)
para("Generated once, after all chunks, from the consolidated state. It lists, per "
     "batch, the successful and failed tickets, and a final summary of fully-completed "
     "tickets vs. those needing manual review. Batch 8 only processes the tickets that "
     "Batch 7 still showed pending in ZSD_DEL_TICKETS.")
code(
    "===================================================\n"
    "  CONSOLIDATED REPORT - ALL CHUNKS\n"
    "===================================================\n"
    "  BATCH 6 - ZCMR Orders     Successful: ...   Failed: ...\n"
    "  BATCH 7 - Final ZSD       Successful: ...   Failed: ...\n"
    "  BATCH 8 - Order Cancel    Successful: ...   Failed: ...\n"
    "  ---------------------------------------------------\n"
    "  100% completed:  N/total      Review manually:  ...\n"
    "==================================================="
)

doc.add_paragraph().add_run("─" * 60).font.color.rgb = ACCENT
foot = doc.add_paragraph()
r = foot.add_run("Validated in production (June 2026): a 200-ticket batch was "
                 "decommissioned end-to-end successfully with no manual intervention.")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = GREEN

doc.save(str(OUT))
print(f"OK -> {OUT}")
