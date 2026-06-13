"""
Genera un documento Word (case study) que responde la pregunta de entrevista:
"An example of a complex AI Automation that you've built. Provide details of
what you built and the outcome it created for the client."

Basado en la arquitectura real del Proyecto DIL y las mejoras implementadas.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parent.parent / "Documentation" / "AI Automation Case Study - SAP DIL Pipeline.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)   # azul corporativo
GREY = RGBColor(0x59, 0x59, 0x59)

doc = Document()

# --- Estilos base ---
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT
    return h


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            r = p.add_run(it[0] + ": ")
            r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def spacer():
    doc.add_paragraph()


# ============================ PORTADA ============================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("Complex Intelligent Automation — Case Study")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = ACCENT

sub = doc.add_paragraph()
r = sub.add_run("Autonomous SAP Ticket-Decommissioning Pipeline (“DIL”)")
r.italic = True
r.font.size = Pt(13)
r.font.color.rgb = GREY

q = doc.add_paragraph()
r = q.add_run("Interview prompt: “An example of a complex AI automation that you've "
              "built. Provide details of what you built and the outcome it created "
              "for the client.”")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = GREY

doc.add_paragraph().add_run("─" * 60).font.color.rgb = ACCENT

# ============================ TL;DR ============================
heading("Executive summary", 1)
doc.add_paragraph(
    "I designed and built an autonomous automation that replaces a slow, manual, "
    "high-risk back-office process in SAP ECC (production). For each transport "
    "ticket, a strict chain of dependent documents must be cancelled in SAP in the "
    "exact right order — billing documents, shipment costs, shipment numbers, "
    "goods-issue reversal, delivery/BOL, and finally the sales or purchase order. "
    "Done by hand this takes minutes per ticket across eight SAP transactions, is "
    "error-prone, and — critically — a single wrong deletion can destroy "
    "records that must be preserved."
)
doc.add_paragraph(
    "The bot drives the real SAP GUI through scripting, runs three parallel SAP "
    "sessions, processes tickets in resilient batches, verifies every step against "
    "SAP itself, recovers automatically from interruptions, and produces an "
    "auditable report. The decision-making core is what makes it “intelligent” "
    "rather than a brittle macro: it inspects each order's real state and decides, "
    "per line, what is safe to delete."
)
p = doc.add_paragraph()
r = p.add_run("Validated in production: a 200-ticket batch ran end-to-end "
              "successfully, decommissioning every targeted ticket and its document "
              "chain with no manual intervention.")
r.bold = True

# ============================ PROBLEM ============================
heading("The business problem", 1)
bullets([
    ("Manual & repetitive", "Hundreds of tickets per cycle, each requiring the same "
     "multi-step teardown across eight different SAP transactions."),
    ("Mandatory ordering", "SAP blocks later steps until earlier documents are removed "
     "— the sequence is not optional, it is enforced by the ERP."),
    ("High risk", "Operations run on PRODUCTION SAP and some are irreversible "
     "(e.g., reversing a goods issue). A mistake is costly and hard to undo."),
    ("Data-integrity hazard", "A single sales order can contain lines that must be "
     "deleted mixed with lines that must be kept. A naive “delete the order” "
     "bot would destroy data that the business needs to retain."),
    ("No audit trail", "The manual process left no consistent record of what was "
     "deleted, what failed, and why."),
])

# ============================ WHAT I BUILT ============================
heading("What I built", 1)
doc.add_paragraph(
    "An end-to-end Python automation that orchestrates the full decommissioning "
    "chain. It reads the target tickets from a spreadsheet, logs into SAP, and runs "
    "an eight-stage pipeline in the dependency order the ERP requires:"
)

table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].paragraphs[0].add_run("Stage").bold = True
hdr[1].paragraphs[0].add_run("SAP transaction").bold = True
hdr[2].paragraphs[0].add_run("What it does").bold = True
rows = [
    ("1 · Billing Documents", "VF11", "Bulk-reverse the invoices tied to each ticket"),
    ("2 · Shipment Cost", "VI05", "Unlock (if needed) and delete the freight cost docs"),
    ("3 · Shipment Number", "VT02N", "Un-assign deliveries and delete the shipment"),
    ("4 · Reverse PGI", "VL09", "Bulk-reverse the post goods issue (only where required)"),
    ("5 · BOL / Delivery", "VL06F", "Delete the outbound/replenishment delivery"),
    ("6 · Orders", "ZCMR → VA02 / ME22N", "Delete the sales / purchase orders"),
    ("7 · Final verification", "ZSD_DEL_TICKETS", "Cross-check against SAP ground truth"),
    ("8 · Order cancellation", "VA02 / ME22N", "Retry whatever step 7 still shows pending"),
]
for a, b, c in rows:
    cells = table.add_row().cells
    cells[0].paragraphs[0].add_run(a)
    cells[1].paragraphs[0].add_run(b)
    cells[2].paragraphs[0].add_run(c)
spacer()

# ============================ HOW IT WORKS ============================
heading("How it works — architecture", 1)
bullets([
    ("Three parallel SAP sessions", "One session reads data and verifies, one performs "
     "deletions, one handles orders. Separating concerns lets the bot verify while it "
     "operates and avoids the sessions interfering with each other."),
    ("Read once, reuse everywhere", "A single VL06F query loads every document number "
     "for the whole batch into memory; all stages read from that snapshot instead of "
     "re-querying SAP — a major performance optimization."),
    ("Chunked execution", "Tickets are processed in groups of 100 to stay within SAP's "
     "work-process timeouts and lock limits, and to contain the blast radius of any "
     "failure."),
    ("Self-healing checkpoints", "Progress is written to disk atomically after each "
     "chunk. If the run is interrupted (timeout, crash, reboot), it resumes exactly "
     "where it left off and skips completed work — fully idempotent."),
    ("Verify against ground truth", "After every stage the bot re-reads SAP to confirm "
     "the document is actually gone. A ticket is only “done” when SAP itself "
     "says so; failures are isolated so one bad ticket never stops the batch."),
    ("Observability", "Every run is mirrored to a timestamped UTF-8 log, giving a full, "
     "auditable trace of every decision and SAP message."),
])

# ============================ THE INTELLIGENT CORE ============================
heading("The hardest part: an order-deletion engine that decides, not assumes", 1)
doc.add_paragraph(
    "The defining challenge — and where the automation becomes genuinely "
    "“intelligent” — was order deletion. In this SAP system a single order "
    "frequently mixes lines that must be deleted with lines that must be kept. A bot "
    "that simply deleted the whole order would silently destroy valid business data."
)
doc.add_paragraph("I built a “verify-then-decide” engine. For each order it:")
bullets([
    "Opens the order once and reads its actual line items;",
    "Matches every line against the exact set of in-scope tickets;",
    ("Decides by evidence, not assumption", "delete the WHOLE order only if every line "
     "is in scope; if the order is mixed, delete ONLY the in-scope lines and preserve "
     "the rest; if the lines are already gone, treat it as success (idempotent)."),
    ("Fails safe", "if it cannot positively classify a line, it stops rather than "
     "guessing — it will never delete data it isn't sure about."),
])
doc.add_paragraph(
    "Around this core I added order-level de-duplication (dozens of tickets often map "
    "to the same order, so each order is processed once and the result is propagated to "
    "all of its tickets) and resilient grid-column matching, because SAP exposes "
    "different technical column names depending on the active screen layout."
)

# ============================ WHY COMPLEX ============================
heading("Why this is a complex automation", 1)
bullets([
    ("Driving a live, stateful, asynchronous UI", "no clean API — the bot scripts "
     "the real SAP GUI, waits on async backend processing, and handles dynamic popups "
     "and shifting screen layouts."),
    ("Eight interdependent transactions", "with a mandatory order enforced by the ERP."),
    ("Correctness under partial failure", "verification, isolation, and a final "
     "ground-truth reconciliation step that corrects any false “success.”"),
    ("Safety on irreversible, production operations", "idempotency, fail-safe defaults, "
     "and never destroying out-of-scope data."),
    ("Resilience", "resumable checkpoints, chunking, and full observability."),
])

# ============================ OUTCOME ============================
heading("Outcome for the client", 1)
bullets([
    ("Proven in production", "a 200-ticket batch was decommissioned end-to-end "
     "successfully, with no manual steps."),
    ("Massive time savings", "a process that took minutes per ticket by hand now runs "
     "unattended; full batches of 150–240+ tickets complete in roughly 15–20 "
     "minutes while the operator does other work."),
    ("Eliminated data-loss risk", "the verify-then-decide engine guarantees that lines "
     "which must be kept are never deleted — something the manual process could not "
     "reliably promise."),
    ("Reliability & recoverability", "interrupted runs resume automatically; re-running "
     "is safe and never double-deletes or errors on already-completed work."),
    ("Auditability", "every run produces a per-ticket, per-stage report plus a complete "
     "log — turning an opaque manual task into a traceable, reviewable process."),
])

# ============================ HOW BUILT ============================
heading("How it was built", 1)
doc.add_paragraph(
    "The system is written in Python and built with an AI-assisted, agentic development "
    "workflow. A recurring pattern proved especially powerful: instrument first, then "
    "fix. When deletions failed, the bot was extended to dump SAP's real screen state "
    "(column names, status-bar messages, menu structure) into the logs, so the true "
    "root cause could be diagnosed from a single production run instead of guesswork. "
    "That instrumentation-driven loop is how subtle, high-stakes issues — such as a "
    "row-selection timing bug and SAP's “document is not in the database” "
    "idempotency case — were found and resolved quickly and safely."
)

doc.add_paragraph().add_run("─" * 60).font.color.rgb = ACCENT
foot = doc.add_paragraph()
r = foot.add_run("Tech stack: Python · SAP GUI Scripting (COM) · openpyxl · "
                 "atomic JSON checkpointing · structured logging.")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = GREY

doc.save(str(OUT))
print(f"OK -> {OUT}")
